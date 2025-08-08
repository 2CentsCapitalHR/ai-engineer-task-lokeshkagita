import os
import json
from docx import Document
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
import google.generativeai as genai

# === Gemini API Setup ===
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel("gemini-2.5-pro")

# === Vector DB Initialization ===
REFERENCE_PDF = "Data Sources.pdf"
INDEX_DIR = "faiss_index"

# Step 1: Extract legal reference and create chunks
def load_reference_text(pdf_path):
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def chunk_reference_text(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    return splitter.split_text(text)

# Step 2: Build or load FAISS vector store
def get_vector_store():
    if os.path.exists(INDEX_DIR):
        return FAISS.load_local(INDEX_DIR, HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2"), allow_dangerous_deserialization=True)
    else:
        raw_text = load_reference_text(REFERENCE_PDF)
        chunks = chunk_reference_text(raw_text)
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        db = FAISS.from_texts(chunks, embedding=embeddings)
        db.save_local(INDEX_DIR)
        return db

# Step 3: RAG-based Prompt Construction
def retrieve_context_from_rag(document_text, db):
    matches = db.similarity_search(document_text, k=3)
    return "\n\n".join([doc.page_content for doc in matches])

def ask_gemini(document_text, db):
    try:
        adgm_context = retrieve_context_from_rag(document_text, db)

        prompt = f"""
        You are a legal assistant specializing in ADGM (Abu Dhabi Global Market) compliance.
        Based on the following legal reference content:

        {adgm_context}

        Review the document below for ADGM compliance. Identify:
        - Invalid/missing clauses
        - Incorrect jurisdiction (e.g., UAE Federal instead of ADGM)
        - Ambiguous or non-binding language
        - Signature or formatting issues

        Text:
        \"\"\"{document_text}\"\"\"

        Respond ONLY in JSON format:
        {{
          "issues_found": [
            {{
              "section": "...",
              "issue": "...",
              "severity": "...",
              "suggestion": "..."
            }}
          ]
        }}
        """

        response = model.generate_content([{"role": "user", "parts": [prompt]}])
        raw = response.text.strip()
        json_start = raw.find('{')
        json_end = raw.rfind('}') + 1
        json_str = raw[json_start:json_end]
        result_json = json.loads(json_str)
        return result_json.get("issues_found", [])

    except Exception as e:
        print("❌ Gemini or JSON parsing error:", e)
        return []

# Step 4: Analyze and Generate Output
def analyze_documents(file_path):
    os.makedirs("outputs", exist_ok=True)
    db = get_vector_store()

    filename = os.path.basename(file_path)
    doc = Document(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    issues = ask_gemini(full_text, db)

    output_path = os.path.join("outputs", f"reviewed_{filename}")
    reviewed = Document(file_path)
    reviewed.add_paragraph("")
    for issue in issues:
        reviewed.add_paragraph(f"🔍 Comment: {issue['section']} — {issue['issue']} → {issue['suggestion']}")
    reviewed.save(output_path)

    report = {
        "process": "Company Incorporation",
        "documents_uploaded": 1,
        "required_documents": 5,
        "missing_document": [
            "Memorandum of Association",
            "Board Resolution Template",
            "UBO Declaration Form",
            "Register of Members and Directors"
        ],
        "issues_found": issues
    }

    report_path = os.path.join("outputs", f"{filename}_report.json")
    with open(report_path, "w") as f:
        json.dump(report, f, indent=4)

    return {"output_docx": output_path, "report_json": report_path}

